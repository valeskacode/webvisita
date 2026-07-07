# -*- coding: utf-8 -*-
"""
Funciones compartidas estilos, lectura de Excel, autoguardado/recuperación de avance, historial y generación de reportes (Word y PDF).
"""
import io
import json
import os
import re
import shutil
from datetime import datetime

import pandas as pd
import streamlit as st

from PIL import Image #importa imagen

try:
    from zoneinfo import ZoneInfo
    TZ_PERU = ZoneInfo("America/Lima")
except Exception:  # pragma: no cover - fallback si no hay tzdata
    TZ_PERU = None


def ahora_peru() -> datetime:
    #Devuelve la fecha/hora ACTUAL
    if TZ_PERU is not None:
        return datetime.now(TZ_PERU)
    return datetime.now()


# rutas
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, "data")
DRAFTS_DIR = os.path.join(DATA_DIR, "drafts")
FOTOS_DIR = os.path.join(DRAFTS_DIR, "fotos")
HISTORIAL_PATH = os.path.join(DATA_DIR, "historial_visitas.xlsx")

#almacenamiento de pdf y word.
REPORTES_DIR = os.environ.get(
    "VISITA_APP_REPORTES_DIR", os.path.join(DATA_DIR, "reportes_generados")
)

# ruta de carpeta
for _d in (DATA_DIR, DRAFTS_DIR, FOTOS_DIR, REPORTES_DIR):
    try:
        os.makedirs(_d, exist_ok=True)
    except Exception as _e:  # ruta inválida (ej. una URL) u otro problema de disco
        if _d is REPORTES_DIR:
            REPORTES_DIR = os.path.join(DATA_DIR, "reportes_generados")
            os.makedirs(REPORTES_DIR, exist_ok=True)

EXCEL_SHEET_NAME = "MUESTRA_FINAL"

EXCEL_COLUMNS = [
    "RECNO", "PEPAIS", "PETDOC", "DOCPEN", "CODCLI", "BCEMP", "BCSUC", "BCMDA",
    "BCPAP", "BCCTA", "BCOPER", "BCSBOP", "BCTOP", "BCMOD", "CODCRE", "REGION",
    "ZONA", "AGENCIA", "CLIENTE", "DIRECCION_DOM", "DISTRITO_DOM",
    "PROVINCIA_DOM", "DEPARTAMENTO_DOM", "DIRECCION_NEG", "DISTRITO_NEG",
    "PROVINCIA_NEG", "DEPARTAMENTO_NEG", "ACTIVIDAD_ECON", "ANALISTA",
    "PRODUCTO_CAJA", "SALDO_MN", "SALDO_VIGE", "SALDO_REFI", "SALDO_VENC",
    "SALDO_JUDI", "MORA_CONT", "TIPO_SBS", "FECDES", "IMPDESEMB_MN",
    "COD_MODULO", "MODULO", "COD_TIPO_OPERACION", "TIPO_OPERACION",
    "ANALISTA_EVAL", "USUARIO_APROB", "USUARIO_DESEM", "FECHA_EVAL",
    "DIAS_ATRASO", "ESTADO_CREDITO", "ATRANT_1M", "ATRANT_2M", "ATRANT_3M",
    "ATRANT_4M", "ATRANT_5M", "ATRANT_6M", "TIPO_SOLI", "NUMERO_CUOTAS",
    "CUOTAS_PAGADAS", "TIPO", "SEGMENTACION_MYPE", "CATEG_RESULTANTE",
    "CATEG_RESULTANTE_SINALIN", "CUENTA_AVAL", "FECHA_UTLPAGO", "UAI_IND",
    "ESTRATO", "TIPO_EXPEDIENTE",
]

# cliente encontrado
COLUMNAS_OPCIONALES_CONTACTO = [
    "TELEFONO",           # Ej: número de celular del cliente
    "EMAIL",               # Ej: correo de contacto
    "LIMITE_CREDITO_MN",   # Ej: línea de crédito aprobada (no es el saldo actual)
]
EXCEL_COLUMNS = EXCEL_COLUMNS + [
    c for c in COLUMNAS_OPCIONALES_CONTACTO if c not in EXCEL_COLUMNS
]

CLIENTE_VISITADO_OPCIONES = [
    "1. Cliente con actividad laboral y/o económica vigente",
    "2. Cliente con situación desmejorada",
    "3. Cliente ya no labora y/o no realiza la actividad económica",
    "4. Cliente no ubicado",
]

CRITERIOS_DEF = {
    "Indicio de dolo o fraude en la evaluación de créditos": [
        "Documentos con enmendaduras",
        "Documentos con datos inconsistentes",
        "Documentos sin datos del cliente",
        "Documentos sin firmas o que no coinciden",
        "Documentos duplicados en más de un cliente",
    ],
    "Evaluaciones deficientes o con sustento insuficiente": [
        "No se evidenció sustento de actividad económica",
        "No se evidenció sustento de ingresos",
        "No se evidenció sustento de activos representativos",
        "Se omitió al cónyuge",
    ],
    "Créditos reprogramados y refinanciados": [
        "Reprogramado",
        "Refinanciado",
    ],
    "Clientes con créditos con calificación diferente a normal a la fecha de revisión": [
        "Calificación diferente a normal",
    ],
}


# estilos
def load_css(path):
    """Inyecta un archivo CSS dentro de la app de Streamlit."""
    full_path = os.path.join(BASE_DIR, path) if not os.path.isabs(path) else path
    try:
        with open(full_path, "r", encoding="utf-8") as f:
            css = f.read()
        st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)
    except FileNotFoundError:
        st.warning(f"No se encontró el archivo de estilos: {full_path}")


# helpers de datos
def safe_str(v, default=""):
    if v is None:
        return default
    try:
        if pd.isna(v):
            return default
    except Exception:
        pass
    s = str(v).strip()
    return default if s.lower() in ("nan", "none") else s


def safe_float(v, default=0.0):
    try:
        f = float(str(v).replace(",", "").strip())
        if pd.isna(f):
            return default
        return f
    except Exception:
        return default


def fmt_money(v):
    return f"S/ {safe_float(v):,.2f}"


def slug(texto):
    """Convierte un texto en algo seguro para usar como nombre de archivo
    o clave de widget, quitando tildes correctamente en vez de dejarlas
    como guiones bajos sueltos (ej. 'Pérez' -> 'Perez', no 'P_rez')."""
    import unicodedata
    texto = safe_str(texto, "sin_dato")
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    texto = re.sub(r"[^A-Za-z0-9_\-]+", "_", texto)
    texto = re.sub(r"_+", "_", texto)
    return texto.strip("_") or "sin_dato"


# --------------------------------------------------------------------------
#  HELPERS PARA LA TARJETA 
# --------------------------------------------------------------------------
def solo_digitos(texto):
    """Deja solo los dígitos de un texto. Útil para comparar DNIs aunque
    el usuario escriba espacios, puntos o el nombre junto al número."""
    return re.sub(r"\D", "", safe_str(texto))


def iniciales(nombre):
    #Devuelve 2 iniciales en mayúscula para el avatar circular
    partes = [p for p in safe_str(nombre).split() if p]
    if not partes:
        return "?"
    if len(partes) == 1:
        return partes[0][:2].upper()
    return (partes[0][0] + partes[1][0]).upper()


def clase_calificacion(calif):
    #color clasificcacion
    c = safe_str(calif).strip().upper()
    if c.startswith("A"):
        return "chip-calif-ok"
    if c.startswith("B"):
        return "chip-calif-warn"
    if c:
        return "chip-calif-bad"
    return "chip-calif-na"


def clientes_similares(df, fila_actual, max_resultados=3):
    """Busca otros clientes con nombre o DNI parecido al de `fila_actual`,
    para ayudar al usuario a distinguir homónimos o detectar posibles
    duplicados (ver criterio "Documentos duplicados en más de un cliente"
    en CRITERIOS_DEF). 
    """
    nombre_actual = safe_str(fila_actual.get("CLIENTE")).strip().lower()
    dni_actual = solo_digitos(fila_actual.get("DOCPEN"))
    palabras_actual = set(nombre_actual.split())

    candidatos = []
    for _, row in df.iterrows():
        nombre_row = safe_str(row.get("CLIENTE")).strip().lower()
        dni_row = solo_digitos(row.get("DOCPEN"))
        if dni_row == dni_actual and nombre_row == nombre_actual:
            continue  # es el mismo cliente, no un "similar"
        palabras_row = set(nombre_row.split())
        comparten_nombre = len(palabras_actual & palabras_row) >= 2
        comparten_dni = len(dni_actual) >= 6 and dni_row[:6] == dni_actual[:6]
        if comparten_nombre or comparten_dni:
            candidatos.append(row)
        if len(candidatos) >= max_resultados:
            break
    return pd.DataFrame(candidatos) if candidatos else pd.DataFrame(columns=df.columns)

# LECTURA DEL EXCEL 
@st.cache_data(show_spinner="Procesando archivo Excel...")
def cargar_excel(file_bytes: bytes):
    bio = io.BytesIO(file_bytes)
    hoja_usada = EXCEL_SHEET_NAME
    try:
        df = pd.read_excel(bio, sheet_name=EXCEL_SHEET_NAME, dtype=str)
    except ValueError:
        bio.seek(0)
        xls = pd.ExcelFile(bio)
        hoja_usada = xls.sheet_names[0]
        df = pd.read_excel(bio, sheet_name=hoja_usada, dtype=str)

    df.columns = [str(c).strip().upper() for c in df.columns]
    # Compatibilidad con archivos antiguos que usaban "PENDOC" en vez de "DOCPEN"
    if "PENDOC" in df.columns and "DOCPEN" not in df.columns:
        df = df.rename(columns={"PENDOC": "DOCPEN"})
    df = df.fillna("")
    faltantes = [c for c in EXCEL_COLUMNS if c not in df.columns]
    return df, hoja_usada, faltantes


# autocargado proceso
INGRESOS_KEYS = [
    "ingreso_principal", "otros_ingresos",
    "op_alquiler", "op_servicios", "op_transporte", "op_mercaderia", "op_publicidad", "op_otros",
    "fam_alimentacion", "fam_vivienda", "fam_servicios", "fam_educacion", "fam_salud", "fam_otros",
]


def _draft_path(usuario, dni):
    return os.path.join(DRAFTS_DIR, f"{slug(usuario)}__{slug(dni)}.json")


def hay_borrador(usuario, dni):
    return os.path.exists(_draft_path(usuario, dni))


def guardar_borrador(usuario, dni, cliente_actual):
    """Guarda el avance actual de la sesión a disco (foto incluida)."""
    
    if not usuario or not dni:
        return
    data = {"cliente_actual": cliente_actual, "guardado_en": ahora_peru().isoformat()}

    data["criterios"] = {k: v for k, v in st.session_state.items() if k.startswith("chk_")}
    data["_criterios_snapshot"] = data["criterios"]  # copia explícita para la vista Reporte
    data["calif_revision"] = st.session_state.get("calif_revision", "")
    data["ingresos_gastos"] = {k: st.session_state.get(k, 0.0) for k in INGRESOS_KEYS}
    data["garantias"] = st.session_state.get("garantias", [])
    data["rcc"] = st.session_state.get("rcc", [])
    data["cliente_visitado"] = st.session_state.get("cliente_visitado", "")
    data["observacion_criterio"] = st.session_state.get("observacion_criterio", "")

    visitas_serializables = {}
    for clave, v in st.session_state.get("visitas", {}).items():
        v = dict(v)
        foto_bytes = v.pop("foto_bytes", None)
        if foto_bytes:
            foto_path = os.path.join(FOTOS_DIR, f"{slug(usuario)}__{slug(dni)}__{clave}.jpg")
            with open(foto_path, "wb") as f:
                f.write(foto_bytes)
            v["foto_path"] = foto_path
        visitas_serializables[clave] = v
    data["visitas"] = visitas_serializables

    with open(_draft_path(usuario, dni), "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, default=str)


def cargar_borrador(usuario, dni):
    """Carga un avance guardado previamente de vuelta a session_state."""
    path = _draft_path(usuario, dni)
    if not os.path.exists(path):
        return False
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    for k, v in data.get("criterios", {}).items():
        st.session_state[k] = v
    st.session_state["_criterios_snapshot"] = data.get("_criterios_snapshot", data.get("criterios", {}))
    st.session_state["calif_revision"] = data.get("calif_revision", "")

    for k, v in data.get("ingresos_gastos", {}).items():
        st.session_state[k] = v

    st.session_state["garantias"] = data.get("garantias", [])
    st.session_state["rcc"] = data.get("rcc", [])
    st.session_state["cliente_visitado"] = data.get("cliente_visitado", "")
    st.session_state["observacion_criterio"] = data.get("observacion_criterio", "")

    visitas = data.get("visitas", {})
    for clave, v in visitas.items():
        foto_path = v.pop("foto_path", None)
        if foto_path and os.path.exists(foto_path):
            with open(foto_path, "rb") as imgf:
                v["foto_bytes"] = imgf.read()
    st.session_state["visitas"] = visitas
    return True


def borrar_borrador(usuario, dni):
    path = _draft_path(usuario, dni)
    if os.path.exists(path):
        os.remove(path)
    for clave in ("domicilio", "negocio", "aval"):
        foto_path = os.path.join(FOTOS_DIR, f"{slug(usuario)}__{slug(dni)}__{clave}.jpg")
        if os.path.exists(foto_path):
            os.remove(foto_path)


# historial
HISTORIAL_COLUMNS = [
    "Usuario_Auditor", "Agencia", "Cliente", "DNI", "Cuenta",
    "N_Visita_Agencia", "N_Visita_General",
    "ClienteVisitado", "Fecha", "Hora",
    "TipoArchivo", "NombreArchivo", "RutaGuardado", "CriteriosSeleccionados",
    "NumeroOperacion", "Modulo", "AnalistaVigente", "AnalistaEvaluador",
]


def leer_historial():
    if not os.path.exists(HISTORIAL_PATH):
        return pd.DataFrame(columns=HISTORIAL_COLUMNS)
    return pd.read_excel(HISTORIAL_PATH)


def _contar_visitas_previas(agencia):
    # registro por agencia
    hist = leer_historial()
    n_general = len(hist)
    if "Agencia" in hist.columns and agencia:
        n_agencia = int((hist["Agencia"].astype(str).str.strip() == str(agencia).strip()).sum())
    else:
        n_agencia = 0
    return n_agencia, n_general


def registrar_historial(usuario, cliente_actual, tipo_archivo, nombre_archivo,
                         criterios_texto, cliente_visitado="", ruta_guardado=""):
    """Agrega una fila al historial general (data/historial_visitas.xlsx).

    Calcula automáticamente N_Visita_Agencia (numerador dentro de esa
    agencia) y N_Visita_General (numerador global, todas las agencias).
    """
    import openpyxl

    agencia = safe_str(cliente_actual.get("AGENCIA"))
    n_agencia_prev, n_general_prev = _contar_visitas_previas(agencia)

    ahora = ahora_peru()
    fila = [
        usuario, agencia,
        safe_str(cliente_actual.get("CLIENTE")),
        safe_str(cliente_actual.get("DOCPEN")),
        safe_str(cliente_actual.get("BCCTA")),
        n_agencia_prev + 1, n_general_prev + 1,
        cliente_visitado,
        ahora.strftime("%d/%m/%Y"), ahora.strftime("%H:%M:%S"),
        tipo_archivo, nombre_archivo, ruta_guardado, criterios_texto,
        safe_str(cliente_actual.get("BCOPER")),
        safe_str(cliente_actual.get("MODULO")),
        safe_str(cliente_actual.get("ANALISTA")),
        safe_str(cliente_actual.get("ANALISTA_EVAL")),
    ]
    if os.path.exists(HISTORIAL_PATH):
        wb = openpyxl.load_workbook(HISTORIAL_PATH)
        ws = wb.active
        encabezado_actual = [c.value for c in ws[1]]
        if encabezado_actual != HISTORIAL_COLUMNS:
            # Migra archivos de historial de una versión anterior de la app
            # sin perder las filas ya registradas.
            ws.delete_rows(1)
            ws.insert_rows(1)
            for i, col in enumerate(HISTORIAL_COLUMNS, start=1):
                ws.cell(row=1, column=i, value=col)
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Historial"
        ws.append(HISTORIAL_COLUMNS)
    ws.append(fila)
    wb.save(HISTORIAL_PATH)
    return n_agencia_prev + 1, n_general_prev + 1


def reporte_consolidado_por_agencia():
    """Cuántos clientes distintos se visitó por agencia y cuántos
    reportes en total, a partir de todo lo guardado en el historial."""
    hist = leer_historial()
    if hist.empty or "Agencia" not in hist.columns:
        return pd.DataFrame(columns=["Agencia", "Clientes_Visitados", "Reportes_Generados", "Ultima_Visita"])
    resumen = (
        hist.groupby("Agencia")
        .agg(Clientes_Visitados=("DNI", "nunique"),
             Reportes_Generados=("NombreArchivo", "count"),
             Ultima_Visita=("Fecha", "max"))
        .reset_index()
        .sort_values("Clientes_Visitados", ascending=False)
    )
    return resumen


def reporte_consolidado_por_cliente(agencia=None):
    """Detalle por cliente: cuántas visitas/reportes tiene cada uno,
    opcionalmente filtrado por agencia."""
    hist = leer_historial()
    if hist.empty:
        return pd.DataFrame(columns=["Agencia", "Cliente", "DNI", "Cuenta", "Reportes_Generados", "Ultima_Visita"])
    if agencia and "Agencia" in hist.columns:
        hist = hist[hist["Agencia"].astype(str).str.strip() == str(agencia).strip()]
    resumen = (
        hist.groupby(["Agencia", "Cliente", "DNI", "Cuenta"], dropna=False)
        .agg(Reportes_Generados=("NombreArchivo", "count"), Ultima_Visita=("Fecha", "max"))
        .reset_index()
        .sort_values("Reportes_Generados", ascending=False)
    )
    return resumen


def reporte_anexo07(agencia=None):
    """Detalle de visitas en el formato oficial 'Anexo 07 — Resultado de
    visitas' (Gerencia de Auditoría Interna): una fila por cliente
    visitado, con cuenta, número de operación, módulo, analista vigente,
    analista evaluador y el resultado/observación de la visita.

    Si el mismo cliente tiene varias filas en el historial (p.ej. porque
    se descargó Word y PDF, o se generó más de un reporte), se conserva
    solo la más reciente.
    """
    columnas = [
        "Agencia", "Cuenta", "NumeroOperacion", "Cliente", "Modulo",
        "AnalistaVigente", "AnalistaEvaluador", "ClienteVisitado",
        "_orden",
    ]
    hist = leer_historial()
    if hist.empty:
        return pd.DataFrame(columns=columnas[:-1])

    if agencia and "Agencia" in hist.columns:
        hist = hist[hist["Agencia"].astype(str).str.strip() == str(agencia).strip()]

    if hist.empty:
        return pd.DataFrame(columns=columnas[:-1])

    hist = hist.copy()
    for col in ("NumeroOperacion", "Modulo", "AnalistaVigente", "AnalistaEvaluador"):
        if col not in hist.columns:
            hist[col] = ""
        hist[col] = hist[col].fillna("").astype(str)

    # Orden cronológico para poder quedarnos con el registro más reciente
    # de cada cliente (fecha + hora, en formato dd/mm/aaaa y HH:MM:SS).
    hist["_orden"] = pd.to_datetime(
        hist["Fecha"].astype(str) + " " + hist["Hora"].astype(str),
        format="%d/%m/%Y %H:%M:%S", errors="coerce",
    )

    clave = ["Cuenta", "DNI"] if "DNI" in hist.columns else ["Cuenta"]
    hist = hist.sort_values("_orden").drop_duplicates(subset=clave, keep="last")

    detalle = hist.rename(columns={
        "Cliente": "Cliente",
        "ClienteVisitado": "ClienteVisitado",
    })[["Agencia", "Cuenta", "NumeroOperacion", "Cliente", "Modulo",
        "AnalistaVigente", "AnalistaEvaluador", "ClienteVisitado", "_orden"]]

    detalle = detalle.sort_values("_orden").drop(columns="_orden").reset_index(drop=True)
    return detalle


def to_excel_anexo07(detalle: pd.DataFrame, agencia: str = "") -> bytes:

    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "anexo 07"

    headers = [
        "N°", "Cuenta Cliente", "Número Operación", "Nombre de Cliente",
        "Módulo", "Analista Vigente", "Analista Evaluador",
        "Resultado de la visita", "Cliente visitado",
    ]
    n_cols = len(headers)

    fill_naranja = PatternFill("solid", fgColor="FFC000")
    borde_fino = Side(style="thin", color="1F4E78")
    borde_celda = Border(top=borde_fino, bottom=borde_fino, left=borde_fino, right=borde_fino)
    negrita = Font(bold=True)

    # --- Cabecera institucional ---
    ws["A1"] = "CAJA AREQUIPA"
    ws["A1"].font = Font(bold=True)
    ws.cell(row=1, column=n_cols).value = ""
    ws.cell(row=1, column=n_cols).font = Font(bold=True)

    ws["A2"] = "GERENCIA DE AUDITORÍA INTERNA"
    ws["A2"].font = Font(bold=True)
    ws.cell(row=2, column=n_cols).value = ""
    ws.cell(row=2, column=n_cols).font = Font(bold=True)

    ws["A3"] = "REF: RESULTADO DE VISITAS"
    ws["A3"].font = Font(bold=True)

    ws["A4"] = f"AGENCIA {agencia or '(TODAS)'}"
    ws["A4"].font = Font(bold=True)

    fila_encabezado = 6
    for j, titulo in enumerate(headers, start=1):
        celda = ws.cell(row=fila_encabezado, column=j, value=titulo)
        celda.font = negrita
        celda.fill = fill_naranja
        celda.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        celda.border = borde_celda

    conteo_resultados = {}
    fila_actual = fila_encabezado + 1
    for i, (_, fila) in enumerate(detalle.iterrows(), start=1):
        resultado_visita = "SI"
        cliente_visitado = safe_str(fila.get("ClienteVisitado")) or "-"
        valores = [
            i,
            safe_str(fila.get("Cuenta")),
            safe_str(fila.get("NumeroOperacion")),
            safe_str(fila.get("Cliente")),
            safe_str(fila.get("Modulo")),
            safe_str(fila.get("AnalistaVigente")),
            safe_str(fila.get("AnalistaEvaluador")),
            resultado_visita,
            cliente_visitado,
        ]
        for j, valor in enumerate(valores, start=1):
            celda = ws.cell(row=fila_actual, column=j, value=valor)
            celda.border = borde_celda
            if j == 1 or j == 8:
                celda.alignment = Alignment(horizontal="center")
        conteo_resultados[cliente_visitado] = conteo_resultados.get(cliente_visitado, 0) + 1
        fila_actual += 1

    # --- Resumen 
    fila_actual += 1
    ws.cell(row=fila_actual, column=1, value="Resumen:").font = negrita
    fila_actual += 1
    orden_categorias = [c for c in CLIENTE_VISITADO_OPCIONES if c in conteo_resultados]
    otras_categorias = [c for c in conteo_resultados if c not in CLIENTE_VISITADO_OPCIONES]
    for categoria in orden_categorias + otras_categorias:
        ws.cell(row=fila_actual, column=1,
                 value=f"{categoria} ({conteo_resultados[categoria]})")
        fila_actual += 1

    anchos = [5, 16, 16, 30, 22, 18, 18, 14, 32]
    for j, ancho in enumerate(anchos, start=1):
        ws.column_dimensions[get_column_letter(j)].width = ancho

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def guardar_reporte_en_carpeta(nombre_archivo: str, contenido_bytes: bytes) -> dict:
    
    resultado = {"local": "", "online": "", "error": ""}

    # 1. Copia local (siempre)
    try:
        destino = os.path.join(REPORTES_DIR, nombre_archivo)
        with open(destino, "wb") as f:
            f.write(contenido_bytes)
        resultado["local"] = destino
    except Exception as e:
        resultado["error"] = f"No se pudo guardar localmente: {e}"

    # 2. OneDrive vía Graph API (si está configurado)
    try:
        from utils.onedrive import credenciales_configuradas, subir_reporte
        if credenciales_configuradas():
            ok, url_o_error = subir_reporte(nombre_archivo, contenido_bytes)
            if ok:
                resultado["online"] = url_o_error
            else:
                resultado["error"] += f" | OneDrive: {url_o_error}"
    except Exception as e:
        resultado["error"] += f" | OneDrive (excepción): {e}"

    return resultado


def sincronizar_historial_onedrive():
    """Sube el Excel de historial a OneDrive cada vez que se genera un
    reporte nuevo, para que quede respaldado en la nube."""
    try:
        from utils.onedrive import credenciales_configuradas, subir_historial
        if not credenciales_configuradas():
            return
        if not os.path.exists(HISTORIAL_PATH):
            return
        with open(HISTORIAL_PATH, "rb") as f:
            contenido = f.read()
        subir_historial(contenido)
    except Exception:
        pass  # silencioso — el historial local ya se guardó


# calculos evaluacion
def calcular_resultado(ing):
    total_ingresos = safe_float(ing.get("ingreso_principal")) + safe_float(ing.get("otros_ingresos"))
    gastos_operativos = sum(safe_float(ing.get(k)) for k in [
        "op_alquiler", "op_servicios", "op_transporte", "op_mercaderia", "op_publicidad", "op_otros",
    ])
    gastos_familiares = sum(safe_float(ing.get(k)) for k in [
        "fam_alimentacion", "fam_vivienda", "fam_servicios", "fam_educacion", "fam_salud", "fam_otros",
    ])
    total_gastos = gastos_operativos + gastos_familiares
    utilidad_neta = total_ingresos - total_gastos
    margen = (utilidad_neta / total_ingresos * 100) if total_ingresos else 0.0
    return {
        "total_ingresos": total_ingresos,
        "gastos_operativos": gastos_operativos,
        "gastos_familiares": gastos_familiares,
        "total_gastos": total_gastos,
        "utilidad_neta": utilidad_neta,
        "margen": margen,
    }


def criterios_seleccionados_lista(criterios, calif_revision):
    seleccionados = []
    for categoria, items in CRITERIOS_DEF.items():
        for item in items:
            key = f"chk_{slug(categoria)}_{slug(item)}"
            if criterios.get(key):
                if item == "Calificación diferente a normal" and calif_revision:
                    seleccionados.append(f"{item} ({calif_revision})")
                else:
                    seleccionados.append(item)
    return seleccionados


# ---------------------------------------------------------------------------
# Diseño de la sección "Visita al negocio / laboral / aval / domicilio"
# (compartido entre Word y PDF) — replica el mockup "III. VISITA AL NEGOCIO"
# ---------------------------------------------------------------------------
AZUL_REPORTE = "1B3A5C"
GRIS_BORDE_REPORTE = "D9DEE6"
VERDE_REPORTE = "1E7E34"
VERDE_FONDO_REPORTE = "EAF7EE"
AMBAR_REPORTE = "B45309"
AMBAR_FONDO_REPORTE = "FEF6E7"
ROJO_REPORTE = "C8102E"
ROJO_FONDO_REPORTE = "FDECEC"

TIPOS_VISITA_INFO = {
    "negocio":   ("VISITA AL NEGOCIO", "Información del establecimiento"),
    "laboral":   ("VISITA AL CENTRO LABORAL", "Información del centro laboral"),
    "aval":      ("VISITA AL AVAL", "Información del lugar visitado"),
    "domicilio": ("VISITA AL DOMICILIO", "Información del domicilio"),
}

RESULTADO_VISITA_MAP = {
    "1": dict(icono="✔", color=VERDE_REPORTE, fondo=VERDE_FONDO_REPORTE,
              titulo="Cliente con actividad económica vigente",
              detalle="Se verificó que el negocio se encuentra operando con normalidad."),
    "2": dict(icono="⚠", color=AMBAR_REPORTE, fondo=AMBAR_FONDO_REPORTE,
              titulo="Cliente con situación desmejorada",
              detalle="Se identificaron señales de deterioro en la actividad verificada."),
    "3": dict(icono="⚠", color=ROJO_REPORTE, fondo=ROJO_FONDO_REPORTE,
              titulo="Cliente ya no labora / no realiza la actividad económica",
              detalle="No se pudo verificar actividad económica vigente en el lugar visitado."),
    "4": dict(icono="✖", color=ROJO_REPORTE, fondo=ROJO_FONDO_REPORTE,
              titulo="Cliente no ubicado",
              detalle="No se logró ubicar al cliente en la dirección registrada."),
}

_CHIP_NUM = {"1": "①", "2": "②", "3": "③", "4": "④"}

# Numeración romana dinámica de secciones del reporte (Word y PDF).
# Permite que "Criterio para la visita" empiece en I y que las secciones de
# visita (Negocio/Laboral/Aval/Domicilio) que no tengan datos simplemente no
# aparezcan, sin dejar huecos en la numeración.
_ROMANOS = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII"]


def _romano(n):
    return _ROMANOS[n - 1] if 1 <= n <= len(_ROMANOS) else str(n)


def _resultado_visita_info(cliente_visitado):
    """Mapea la opción elegida en 'Cliente visitado' a icono/color/texto."""
    if not cliente_visitado:
        return None
    num = cliente_visitado.strip().split(".")[0].strip()
    return RESULTADO_VISITA_MAP.get(num)


def _direccion_registrada(d):
    return ", ".join([x for x in [
        d.get("direccion"), d.get("distrito"), d.get("provincia"), d.get("departamento"),
    ] if x]) or "-"


def _partir_comentarios(comentarios):
    comentarios = (comentarios or "").strip()
    if not comentarios:
        return []
    partes = [c.strip() for c in comentarios.replace("\r", "").split("\n") if c.strip()]
    if len(partes) == 1:
        crudo = partes[0]
        trozos = [t.strip().rstrip(".") for t in crudo.split(". ") if t.strip()]
        partes = [t + "." for t in trozos if t]
    return partes


# --------------------------- helpers de bajo nivel (Word) ------------------
def _docx_shade_cell(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _docx_cell_borders(cell, color=GRIS_BORDE_REPORTE, sz=6):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _docx_cell_margins(cell, top=120, bottom=120, left=140, right=140):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _docx_no_table_borders(table):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tblPr.append(borders)


def _docx_clear_cell(cell):
    cell.text = ""
    return cell.paragraphs[0]


def _docx_card_titulo(cell, numero, texto):
    from docx.shared import Pt, RGBColor
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(_CHIP_NUM.get(str(numero), str(numero)))
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor.from_string(AZUL_REPORTE)
    r2 = p.add_run("  " + texto)
    r2.bold = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor.from_string(AZUL_REPORTE)
    return p


def _docx_kv_lines(cell, pairs, label_color="55606E", value_color="1A1A1A"):
    from docx.shared import Pt, RGBColor
    for label, value in pairs:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        r1 = p.add_run(f"{label}:  ")
        r1.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor.from_string(label_color)
        r2 = p.add_run(value if value not in (None, "") else "-")
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor.from_string(value_color)


def _docx_section_banner(doc, numero, titulo):
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _docx_shade_cell(cell, AZUL_REPORTE)
    _docx_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    p = _docx_clear_cell(cell)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{numero}. {titulo}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string("FFFFFF")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _docx_banner_texto(doc, texto_completo):
    """Igual que _docx_section_banner pero recibe el título ya formateado
    completo (p. ej. 'I. Datos del cliente y crédito' o '0.1 Observación'),
    para evitar duplicar puntos en numeraciones tipo '0.1'."""
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _docx_shade_cell(cell, AZUL_REPORTE)
    _docx_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    p = _docx_clear_cell(cell)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(texto_completo)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string("FFFFFF")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _docx_abrir_tarjeta(doc):
    """Crea una tarjeta (tabla de 1 celda con borde) y devuelve la celda lista
    para recibir contenido — reutilizada por todas las secciones del reporte."""
    tabla = doc.add_table(rows=1, cols=1)
    _docx_no_table_borders(tabla)
    c = tabla.cell(0, 0)
    _docx_cell_borders(c)
    _docx_cell_margins(c)
    _docx_clear_cell(c)
    return c


def _docx_cerrar_tarjeta(doc, space_after=10):
    from docx.shared import Pt
    doc.add_paragraph().paragraph_format.space_after = Pt(space_after)


def _docx_seccion_kv(doc, titulo, pairs, dos_columnas=True):
    """Banner + tarjeta(s) de pares clave:valor — para 'Datos del cliente',
    'Ingresos y gastos' o 'Conformidad'. `titulo` ya debe incluir la numeración
    (p. ej. 'I. Datos del cliente y crédito')."""
    _docx_banner_texto(doc, titulo)
    if dos_columnas and len(pairs) > 4:
        mitad = (len(pairs) + 1) // 2
        izquierda, derecha = pairs[:mitad], pairs[mitad:]
        tabla = doc.add_table(rows=1, cols=2)
        _docx_no_table_borders(tabla)
        cl, cr = tabla.cell(0, 0), tabla.cell(0, 1)
        _docx_cell_borders(cl)
        _docx_cell_borders(cr)
        _docx_cell_margins(cl)
        _docx_cell_margins(cr)
        _docx_clear_cell(cl)
        _docx_clear_cell(cr)
        _docx_kv_lines(cl, izquierda)
        _docx_kv_lines(cr, derecha)
    else:
        c = _docx_abrir_tarjeta(doc)
        _docx_kv_lines(c, pairs)
    _docx_cerrar_tarjeta(doc)


def _docx_seccion_lista(doc, titulo, items):
    """Banner + tarjeta con viñetas — para 'Criterio para la visita'."""
    from docx.shared import Pt, Cm
    _docx_banner_texto(doc, titulo)
    c = _docx_abrir_tarjeta(doc)
    for it in items:
        p = c.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run("•  " + str(it))
        r.font.size = Pt(9)
    _docx_cerrar_tarjeta(doc)


def _docx_seccion_parrafo(doc, titulo, texto):
    """Banner + tarjeta con un párrafo de texto libre — para 'Observación'."""
    from docx.shared import Pt
    _docx_banner_texto(doc, titulo)
    c = _docx_abrir_tarjeta(doc)
    p = c.add_paragraph()
    r = p.add_run(texto)
    r.font.size = Pt(9)
    _docx_cerrar_tarjeta(doc)


def _docx_seccion_grupos(doc, titulo, grupos, etiqueta_grupo):
    """Banner + tarjeta con varios sub-grupos de pares clave:valor — para
    'Garantías' y 'Deuda RCC', donde puede haber varios registros."""
    from docx.shared import Pt, RGBColor
    _docx_banner_texto(doc, titulo)
    c = _docx_abrir_tarjeta(doc)
    total = len(grupos)
    for idx, pairs in enumerate(grupos, start=1):
        if total > 1:
            p = c.add_paragraph()
            p.paragraph_format.space_before = Pt(4) if idx > 1 else Pt(0)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{etiqueta_grupo} {idx}")
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor.from_string(AZUL_REPORTE)
        _docx_kv_lines(c, pairs)
    _docx_cerrar_tarjeta(doc)


def add_visita_card_docx(doc, numero_seccion, clave, etiqueta, d, cliente_visitado):
    """Sección con tarjetas al estilo del mockup: banner + 4 tarjetas numeradas
    (Información del lugar, Resultado de la visita, Observaciones, Evidencia)."""
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    titulo_seccion, titulo_info = TIPOS_VISITA_INFO.get(
        clave, (etiqueta.upper(), "Información del lugar visitado")
    )
    if not d:
        # Si no se registró información para esta visita, se omite la sección
        # por completo (no se muestra banner ni advertencia).
        return

    _docx_section_banner(doc, numero_seccion, titulo_seccion)

    resultado = _resultado_visita_info(cliente_visitado) if clave == "negocio" else None

    # Fila 1: Información del lugar  |  Resultado de la visita (solo negocio)
    cols = 2 if resultado else 1
    fila1 = doc.add_table(rows=1, cols=cols)
    _docx_no_table_borders(fila1)

    c_info = fila1.cell(0, 0)
    _docx_cell_borders(c_info)
    _docx_cell_margins(c_info)
    _docx_clear_cell(c_info)
    _docx_card_titulo(c_info, "1", titulo_info)
    _docx_kv_lines(c_info, [
        ("Dirección", d.get("direccion", "-")),
        ("Distrito", d.get("distrito", "-")),
        ("Provincia", d.get("provincia", "-")),
        ("Departamento", d.get("departamento", "-")),
        ("Referencia", d.get("referencia", "-")),
    ])

    if resultado:
        c_res = fila1.cell(0, 1)
        _docx_cell_borders(c_res)
        _docx_shade_cell(c_res, resultado["fondo"])
        _docx_cell_margins(c_res)
        _docx_clear_cell(c_res)
        _docx_card_titulo(c_res, "2", "Resultado de la visita")

        p_icon = c_res.add_paragraph()
        p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_icon.paragraph_format.space_after = Pt(2)
        r = p_icon.add_run(resultado["icono"])
        r.font.size = Pt(26)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(resultado["color"])

        p_lbl = c_res.add_paragraph()
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lbl.paragraph_format.space_after = Pt(2)
        r = p_lbl.add_run("Cliente visitado")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string("55606E")

        p_tit = c_res.add_paragraph()
        p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_tit.paragraph_format.space_after = Pt(6)
        r = p_tit.add_run(resultado["titulo"])
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor.from_string(resultado["color"])

        p_det = c_res.add_paragraph()
        p_det.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_det.paragraph_format.space_after = Pt(0)
        r = p_det.add_run(resultado["detalle"])
        r.font.size = Pt(8.5)
        r.italic = True
        r.font.color.rgb = RGBColor.from_string("444444")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Fila 2: Observaciones del auditor
    numero_obs = "2" if not resultado else "3"
    tabla_obs = doc.add_table(rows=1, cols=1)
    _docx_no_table_borders(tabla_obs)
    c_obs = tabla_obs.cell(0, 0)
    _docx_cell_borders(c_obs)
    _docx_cell_margins(c_obs)
    _docx_clear_cell(c_obs)
    _docx_card_titulo(c_obs, numero_obs, "Observaciones del auditor")

    p = c_obs.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run("Entrevista realizada con:  ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor.from_string("55606E")
    r2 = p.add_run(d.get("entrevista_con") or "-")
    r2.font.size = Pt(9)

    p2 = c_obs.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    r = p2.add_run("Comentarios:")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("55606E")

    partes = _partir_comentarios(d.get("comentarios"))
    if partes:
        for parte in partes:
            pp = c_obs.add_paragraph()
            pp.paragraph_format.space_after = Pt(3)
            pp.paragraph_format.left_indent = Cm(0.3)
            r = pp.add_run("•  " + parte)
            r.font.size = Pt(9)
    else:
        pp = c_obs.add_paragraph()
        r = pp.add_run("Sin comentarios registrados.")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string("888888")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Fila 3: Evidencia de verificación (foto | geolocalización)
    numero_evid = "3" if not resultado else "4"
    fila3 = doc.add_table(rows=1, cols=2)
    _docx_no_table_borders(fila3)

    c_foto = fila3.cell(0, 0)
    _docx_cell_borders(c_foto)
    _docx_cell_margins(c_foto)
    _docx_clear_cell(c_foto)
    _docx_card_titulo(c_foto, numero_evid, "Evidencia de verificación")
    p = c_foto.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Fotografía tomada en la visita")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("55606E")

    if d.get("foto_bytes"):
        try:
            foto_stream = io.BytesIO(d["foto_bytes"])
            img = Image.open(foto_stream)
            buffer_limpio = io.BytesIO()
            img.convert("RGB").save(buffer_limpio, format="PNG")
            buffer_limpio.seek(0)
            c_foto.add_paragraph().add_run().add_picture(buffer_limpio, width=Cm(7.2))
        except Exception as e:
            pp = c_foto.add_paragraph()
            r = pp.add_run("⚠ Error al procesar la imagen de la visita.")
            r.italic = True
            r.font.size = Pt(9)
            print(f"Error técnico al insertar imagen: {e}")
    else:
        pp = c_foto.add_paragraph()
        r = pp.add_run("Sin fotografía registrada.")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string("888888")

    c_geo = fila3.cell(0, 1)
    _docx_cell_borders(c_geo)
    _docx_cell_margins(c_geo)
    _docx_clear_cell(c_geo)
    p = c_geo.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Datos de geolocalización")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("55606E")

    lat, lon = d.get("lat"), d.get("lon")
    coords_txt = f"{lat:.6f}, {lon:.6f}" if lat and lon else "No capturada"
    _docx_kv_lines(c_geo, [
        ("Coordenadas GPS", coords_txt),
        ("Fecha de captura", d.get("fecha", "-")),
        ("Hora de captura", d.get("hora", "-")),
        ("Dirección registrada", _direccion_registrada(d)),
    ])

    if lat and lon:
        p_link = c_geo.add_paragraph()
        p_link.paragraph_format.space_before = Pt(4)
        r = p_link.add_run(f"Ver en Google Maps: https://maps.google.com/?q={lat},{lon}")
        r.font.size = Pt(8)
        r.italic = True
        r.font.color.rgb = RGBColor.from_string("1A5FB4")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Nota final
    p_nota = doc.add_paragraph()
    p_nota.paragraph_format.space_before = Pt(2)
    p_nota.paragraph_format.space_after = Pt(10)
    r = p_nota.add_run(
        "Nota: La información incluida en esta sección ha sido registrada durante la "
        "visita de verificación y constituye evidencia de la evaluación realizada."
    )
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("777777")


# --------------------------- helpers de bajo nivel (PDF) --------------------
_PDF_FUENTES_REGISTRADAS = False


def _pdf_registrar_fuentes():
    """Registra DejaVu Sans (soporta ✔ ⚠ ① etc.) una sola vez por proceso.

    Si las fuentes TTF no están disponibles en el entorno (por ejemplo,
    porque cambió la ruta del sistema o el paquete de fuentes no está
    instalado), se registran alias "DejaVuSans" / "DejaVuSans-Bold" /
    "DejaVuSans-Oblique" apuntando a las fuentes estándar de ReportLab
    (Helvetica). Así el resto del código, que referencia esos nombres de
    fuente en sus ParagraphStyle, sigue funcionando en vez de fallar con
    "Can't map determine family/bold/italic for ...".
    """
    global _PDF_FUENTES_REGISTRADAS
    if _PDF_FUENTES_REGISTRADAS:
        return
    import copy
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    posibles_bases = [
        "/usr/share/fonts/truetype/dejavu/",
        "/usr/share/fonts/dejavu/",
        os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "fonts") + os.sep,
    ]
    archivos = {
        "DejaVuSans": "DejaVuSans.ttf",
        "DejaVuSans-Bold": "DejaVuSans-Bold.ttf",
        "DejaVuSans-Oblique": "DejaVuSans-Oblique.ttf",
    }

    registradas_ok = False
    for base in posibles_bases:
        try:
            if all(os.path.isfile(base + archivo) for archivo in archivos.values()):
                for nombre, archivo in archivos.items():
                    pdfmetrics.registerFont(TTFont(nombre, base + archivo))
                registradas_ok = True
                break
        except Exception:
            continue

    if not registradas_ok:
        # Fallback: alias sobre fuentes estándar (no soportan ✔ ⚠ ① pero
        # evitan que la generación de PDF se caiga por completo).
        equivalentes = {
            "DejaVuSans": "Helvetica",
            "DejaVuSans-Bold": "Helvetica-Bold",
            "DejaVuSans-Oblique": "Helvetica-Oblique",
        }
        for alias, estandar in equivalentes.items():
            try:
                fuente_clon = copy.copy(pdfmetrics.getFont(estandar))
                fuente_clon.fontName = alias
                pdfmetrics.registerFont(fuente_clon)
            except Exception:
                pass
        try:
            # Los alias base14 no activan automáticamente el mapeo bold/italic
            # que usa el parser de Paragraph (a diferencia de TTFont, que sí lo
            # hace). Lo registramos explícitamente para que estilos con
            # fontName="DejaVuSans-Bold"/"DejaVuSans-Oblique" resuelvan bien.
            pdfmetrics.registerFontFamily(
                "DejaVuSans",
                normal="DejaVuSans",
                bold="DejaVuSans-Bold",
                italic="DejaVuSans-Oblique",
                boldItalic="DejaVuSans-Bold",
            )
        except Exception:
            pass

    _PDF_FUENTES_REGISTRADAS = True


def _pdf_estilos_visita():
    from reportlab.lib import colors as _colors
    from reportlab.lib.styles import ParagraphStyle
    azul = _colors.HexColor("#" + AZUL_REPORTE)
    gris_txt = _colors.HexColor("#55606E")
    return {
        "card_title": ParagraphStyle("cardtitle", fontName="DejaVuSans-Bold", fontSize=10.5,
                                      textColor=azul, spaceAfter=6),
        "label": ParagraphStyle("label", fontName="DejaVuSans", fontSize=9,
                                 textColor=_colors.HexColor("#1A1A1A"), leading=13),
        "result_lbl": ParagraphStyle("resultlbl", fontName="DejaVuSans", fontSize=9,
                                      textColor=gris_txt, alignment=1, spaceAfter=2),
        "result_titulo": ParagraphStyle("resulttit", fontName="DejaVuSans-Bold", fontSize=11,
                                         alignment=1, spaceAfter=6, leading=14),
        "result_det": ParagraphStyle("resultdet", fontName="DejaVuSans-Oblique", fontSize=8.5,
                                      textColor=_colors.HexColor("#444444"), alignment=1, leading=11),
        "icon_big": ParagraphStyle("iconbig", fontName="DejaVuSans-Bold", fontSize=28, alignment=1),
        "sub": ParagraphStyle("sub", fontName="DejaVuSans-Bold", fontSize=9,
                               textColor=gris_txt, spaceAfter=4),
        "bullet": ParagraphStyle("bullet", fontName="DejaVuSans", fontSize=9, leftIndent=10,
                                  leading=12.5, spaceAfter=3),
        "italic_muted": ParagraphStyle("italicmuted", fontName="DejaVuSans-Oblique", fontSize=9,
                                        textColor=_colors.HexColor("#888888")),
        "link": ParagraphStyle("link", fontName="DejaVuSans-Oblique", fontSize=8,
                                textColor=_colors.HexColor("#1A5FB4")),
        "nota": ParagraphStyle("nota", fontName="DejaVuSans-Oblique", fontSize=8,
                                textColor=_colors.HexColor("#777777"), leading=11),
        "banner": ParagraphStyle("banner", fontName="DejaVuSans-Bold", fontSize=12.5,
                                  textColor=_colors.white, leading=15),
        "warn_missing": ParagraphStyle("warnmissing", fontName="DejaVuSans-Oblique", fontSize=9.5,
                                        textColor=_colors.HexColor("#" + ROJO_REPORTE)),
        "grupo_titulo": ParagraphStyle("grupotitulo", fontName="DejaVuSans-Bold", fontSize=9.5,
                                        textColor=azul, spaceAfter=4),
    }


def _pdf_banner(numero, titulo, ancho):
    """Barra azul de encabezado, ancho completo — reutilizada por todas las secciones."""
    from reportlab.lib import colors as _colors
    from reportlab.platypus import Paragraph, Table, TableStyle
    est = _pdf_estilos_visita()
    banner = Table([[Paragraph(f"{numero}. {titulo}", est["banner"])]], colWidths=[ancho])
    banner.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), _colors.HexColor("#" + AZUL_REPORTE)),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    return banner


def _pdf_card_wrap(flowables, col_width, fondo=None, borde=None):
    """Tarjeta con borde/relleno — reutilizada por todas las secciones."""
    from reportlab.lib import colors as _colors
    from reportlab.lib.units import cm
    from reportlab.platypus import Table, TableStyle
    if fondo is None:
        fondo = _colors.white
    if borde is None:
        borde = _colors.HexColor("#" + GRIS_BORDE_REPORTE)
    inner = Table([[flowables]], colWidths=[col_width - 0.5 * cm])
    inner.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), fondo),
        ("BOX", (0, 0), (-1, -1), 0.75, borde),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    return inner


def _pdf_banner_texto(texto_completo, ancho):
    """Igual que _pdf_banner pero recibe el título ya formateado completo
    (p. ej. 'I. Datos del cliente y crédito' o '0.1 Observación'), para evitar
    duplicar puntos en numeraciones tipo '0.1'."""
    from reportlab.lib import colors as _colors
    from reportlab.platypus import Paragraph, Table, TableStyle
    est = _pdf_estilos_visita()
    banner = Table([[Paragraph(texto_completo, est["banner"])]], colWidths=[ancho])
    banner.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), _colors.HexColor("#" + AZUL_REPORTE)),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    return banner


def _pdf_kv_paragraphs(pairs, est):
    from reportlab.platypus import Paragraph
    return [Paragraph(f"<b>{lbl}:</b>  {val if val not in (None, '') else '-'}", est["label"])
            for lbl, val in pairs]


def _pdf_seccion_kv(titulo, ancho, pairs, dos_columnas=True):
    """Banner + tarjeta(s) de pares clave:valor — para secciones tipo 'Datos del
    cliente', 'Ingresos y gastos' o 'Conformidad'. `titulo` ya debe incluir la
    numeración (p. ej. 'I. Datos del cliente y crédito')."""
    from reportlab.platypus import Spacer, Table, TableStyle
    _pdf_registrar_fuentes()
    est = _pdf_estilos_visita()
    elems = [_pdf_banner_texto(titulo, ancho), Spacer(1, 8)]
    if dos_columnas and len(pairs) > 4:
        mitad = (len(pairs) + 1) // 2
        izquierda, derecha = pairs[:mitad], pairs[mitad:]
        col_w = ancho / 2
        fila = Table([[_pdf_card_wrap(_pdf_kv_paragraphs(izquierda, est), col_w),
                       _pdf_card_wrap(_pdf_kv_paragraphs(derecha, est), col_w)]],
                     colWidths=[col_w, col_w])
        fila.setStyle(TableStyle([
            ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        elems.append(fila)
    else:
        elems.append(_pdf_card_wrap(_pdf_kv_paragraphs(pairs, est), ancho))
    elems.append(Spacer(1, 12))
    return elems


def _pdf_seccion_lista(titulo, ancho, items):
    """Banner + tarjeta con viñetas — para 'Criterio para la visita'."""
    from reportlab.platypus import Paragraph, Spacer
    _pdf_registrar_fuentes()
    est = _pdf_estilos_visita()
    contenido = [Paragraph("•  " + str(it), est["bullet"]) for it in items]
    return [_pdf_banner_texto(titulo, ancho), Spacer(1, 8),
            _pdf_card_wrap(contenido, ancho), Spacer(1, 12)]


def _pdf_seccion_parrafo(titulo, ancho, texto):
    """Banner + tarjeta con un párrafo de texto libre — para 'Observación'."""
    from reportlab.platypus import Paragraph, Spacer
    _pdf_registrar_fuentes()
    est = _pdf_estilos_visita()
    return [_pdf_banner_texto(titulo, ancho), Spacer(1, 8),
            _pdf_card_wrap([Paragraph(texto, est["label"])], ancho), Spacer(1, 12)]


def _pdf_seccion_grupos(titulo, ancho, grupos, etiqueta_grupo):
    """Banner + tarjeta con varios sub-grupos de pares clave:valor — para
    'Garantías' y 'Deuda RCC', donde puede haber varios registros."""
    from reportlab.platypus import Paragraph, Spacer
    _pdf_registrar_fuentes()
    est = _pdf_estilos_visita()
    contenido = []
    total = len(grupos)
    for idx, pairs in enumerate(grupos, start=1):
        if total > 1:
            contenido.append(Paragraph(f"{etiqueta_grupo} {idx}", est["grupo_titulo"]))
        contenido += _pdf_kv_paragraphs(pairs, est)
        if idx < total:
            contenido.append(Spacer(1, 6))
    return [_pdf_banner_texto(titulo, ancho), Spacer(1, 8),
            _pdf_card_wrap(contenido, ancho), Spacer(1, 12)]


def build_visita_section_pdf(numero_seccion, clave, etiqueta, d, cliente_visitado, ancho_total):
    """Devuelve la lista de flowables reportlab para una sección de visita,
    con el mismo diseño de tarjetas que la versión Word."""
    from reportlab.lib import colors as _colors
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle, Image as RLImage

    _pdf_registrar_fuentes()
    est = _pdf_estilos_visita()

    def kv_paragraphs(pairs):
        return _pdf_kv_paragraphs(pairs, est)

    def card_wrap(flowables, col_width, fondo=_colors.white, borde=_colors.HexColor("#" + GRIS_BORDE_REPORTE)):
        return _pdf_card_wrap(flowables, col_width, fondo=fondo, borde=borde)

    titulo_seccion, titulo_info = TIPOS_VISITA_INFO.get(
        clave, (etiqueta.upper(), "Información del lugar visitado")
    )

    if not d:
        # Si no se registró información para esta visita, se omite la sección
        # por completo (no se muestra banner ni advertencia).
        return []

    elems = [_pdf_banner(numero_seccion, titulo_seccion, ancho_total), Spacer(1, 8)]

    resultado = _resultado_visita_info(cliente_visitado) if clave == "negocio" else None

    if resultado:
        col_w = ancho_total / 2
        info_flow = [Paragraph(f"①  {titulo_info}", est["card_title"])]
        info_flow += kv_paragraphs([
            ("Dirección", d.get("direccion", "-")),
            ("Distrito", d.get("distrito", "-")),
            ("Provincia", d.get("provincia", "-")),
            ("Departamento", d.get("departamento", "-")),
            ("Referencia", d.get("referencia", "-")),
        ])
        res_color = _colors.HexColor("#" + resultado["color"])
        res_flow = [
            Paragraph("②  Resultado de la visita", est["card_title"]),
            Paragraph(resultado["icono"], ParagraphStyle_iconc(est, res_color)),
            Paragraph("Cliente visitado", est["result_lbl"]),
            Paragraph(resultado["titulo"], ParagraphStyle_titc(est, res_color)),
            Paragraph(resultado["detalle"], est["result_det"]),
        ]
        fila1 = Table([[card_wrap(info_flow, col_w),
                         card_wrap(res_flow, col_w, fondo=_colors.HexColor("#" + resultado["fondo"]))]],
                       colWidths=[col_w, col_w])
        fila1.setStyle(TableStyle([
            ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        elems.append(fila1)
        num_obs, num_evid = "③", "④"
    else:
        info_flow = [Paragraph(f"①  {titulo_info}", est["card_title"])]
        info_flow += kv_paragraphs([
            ("Dirección", d.get("direccion", "-")),
            ("Distrito", d.get("distrito", "-")),
            ("Provincia", d.get("provincia", "-")),
            ("Departamento", d.get("departamento", "-")),
            ("Referencia", d.get("referencia", "-")),
        ])
        elems.append(card_wrap(info_flow, ancho_total))
        num_obs, num_evid = "②", "③"

    elems.append(Spacer(1, 8))

    obs_flow = [Paragraph(f"{num_obs}  Observaciones del auditor", est["card_title"]),
                Paragraph(f"<b>Entrevista realizada con:</b>  {d.get('entrevista_con') or '-'}", est["label"]),
                Spacer(1, 4), Paragraph("Comentarios:", est["sub"])]
    partes = _partir_comentarios(d.get("comentarios"))
    if partes:
        for parte in partes:
            obs_flow.append(Paragraph("•  " + parte, est["bullet"]))
    else:
        obs_flow.append(Paragraph("Sin comentarios registrados.", est["italic_muted"]))
    elems.append(card_wrap(obs_flow, ancho_total))
    elems.append(Spacer(1, 8))

    col_w = ancho_total / 2
    foto_flow = [Paragraph(f"{num_evid}  Evidencia de verificación", est["card_title"]),
                 Paragraph("Fotografía tomada en la visita", est["sub"])]
    if d.get("foto_bytes"):
        try:
            img = RLImage(io.BytesIO(d["foto_bytes"]), width=col_w - 1.2 * cm,
                          height=(col_w - 1.2 * cm) * 0.72)
            foto_flow.append(img)
        except Exception:
            foto_flow.append(Paragraph("⚠ Error al procesar la imagen de la visita.", est["italic_muted"]))
    else:
        foto_flow.append(Paragraph("Sin fotografía registrada.", est["italic_muted"]))

    lat, lon = d.get("lat"), d.get("lon")
    coords_txt = f"{lat:.6f}, {lon:.6f}" if lat and lon else "No capturada"
    geo_flow = [Paragraph("Datos de geolocalización", est["sub"])]
    geo_flow += kv_paragraphs([
        ("Coordenadas GPS", coords_txt),
        ("Fecha de captura", d.get("fecha", "-")),
        ("Hora de captura", d.get("hora", "-")),
        ("Dirección registrada", _direccion_registrada(d)),
    ])
    if lat and lon:
        geo_flow.append(Spacer(1, 3))
        geo_flow.append(Paragraph(
            f"Ver en Google Maps: https://maps.google.com/?q={lat},{lon}", est["link"]))

    fila3 = Table([[card_wrap(foto_flow, col_w), card_wrap(geo_flow, col_w)]],
                   colWidths=[col_w, col_w])
    fila3.setStyle(TableStyle([
        ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elems.append(fila3)
    elems.append(Spacer(1, 6))
    elems.append(Paragraph(
        "Nota: La información incluida en esta sección ha sido registrada durante la visita de "
        "verificación y constituye evidencia de la evaluación realizada.", est["nota"]))
    elems.append(Spacer(1, 12))
    return elems


def ParagraphStyle_iconc(est, color):
    from reportlab.lib.styles import ParagraphStyle
    return ParagraphStyle("iconc", parent=est["icon_big"], textColor=color)


def ParagraphStyle_titc(est, color):
    from reportlab.lib.styles import ParagraphStyle
    return ParagraphStyle("titc", parent=est["result_titulo"], textColor=color)


# -reporte word
def generar_word(cliente, criterios_txt, ingresos_calc, ingresos_raw, visitas, garantias, rcc, usuario, cliente_visitado="", observacion_criterio=""):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading("VISITA A CLIENTES", level=0)
    p = doc.add_paragraph("CMAC Caja Arequipa — Unidad de Auditoría Interna")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Auditor: {usuario}  ·  Fecha de visita: {ahora_peru().strftime('%d/%m/%Y %H:%M')} (hora Perú)")

    numero = 1

    if criterios_txt:
        _docx_seccion_lista(doc, f"{_romano(numero)}. Criterio para la visita", criterios_txt)
        if observacion_criterio:
            _docx_seccion_parrafo(doc, f"{_romano(numero)}.1 Observación", observacion_criterio)
        numero += 1
    elif observacion_criterio:
        _docx_seccion_parrafo(doc, f"{_romano(numero)}. Observación", observacion_criterio)
        numero += 1

    _docx_seccion_kv(doc, f"{_romano(numero)}. Datos del cliente y crédito", [
        ("Agencia", safe_str(cliente.get("AGENCIA"))),
        ("DNI/LE Titular", safe_str(cliente.get("DOCPEN"))),
        ("Titular", safe_str(cliente.get("CLIENTE"))),
        ("Cuenta cliente", safe_str(cliente.get("BCCTA"))),
        ("N° de operación", safe_str(cliente.get("BCOPER"))),
        ("Módulo", safe_str(cliente.get("MODULO"))),
        ("Analista vigente", safe_str(cliente.get("ANALISTA"))),
        ("Analista evaluador", safe_str(cliente.get("ANALISTA_EVAL"))),
        ("Auditor (visita)", usuario),
        ("Importe", fmt_money(cliente.get("IMPDESEMB_MN"))),
        ("Saldo capital", fmt_money(cliente.get("SALDO_MN"))),
        ("Tipo de crédito", safe_str(cliente.get("PRODUCTO_CAJA"))),
        ("Tipo SBS", safe_str(cliente.get("TIPO_SBS"))),
        ("Calificación", safe_str(cliente.get("CATEG_RESULTANTE"))),
        ("Rubro", safe_str(cliente.get("ACTIVIDAD_ECON"))),
        ("Último pago", safe_str(cliente.get("FECHA_UTLPAGO"))),
        ("Resultado de la visita / Cliente visitado", cliente_visitado or "-"),
    ])
    numero += 1

    for clave, etiqueta in [("negocio", "Negocio"),
                             ("laboral", "Laboral"),
                             ("aval", "Aval"),
                             ("domicilio", "Domicilio")]:
        d = visitas.get(clave)
        if not d:
            continue  # solo se incluyen las visitas con información registrada
        add_visita_card_docx(doc, _romano(numero), clave, etiqueta, d, cliente_visitado)
        numero += 1

    if garantias:
        _docx_seccion_grupos(doc, f"{_romano(numero)}. Garantías", [list(g.items()) for g in garantias], "Garantía")
        numero += 1

    if rcc:
        _docx_seccion_grupos(doc, f"{_romano(numero)}. Deuda RCC", [list(r.items()) for r in rcc], "Deuda")
        numero += 1

    _docx_seccion_kv(doc, f"{_romano(numero)}. Conformidad", [
        ("Hecho por (Auditor)", usuario), ("Fecha", ahora_peru().strftime("%d/%m/%Y")),
    ], dos_columnas=False)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


#  reporte pdf
def generar_pdf(cliente, criterios_txt, ingresos_calc, ingresos_raw, visitas, garantias, rcc, usuario, cliente_visitado="", observacion_criterio=""):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    _pdf_registrar_fuentes()
    AZUL = colors.HexColor("#" + AZUL_REPORTE)

    buf = io.BytesIO()
    docpdf = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1c", parent=styles["Heading1"], textColor=AZUL, fontSize=15,
                         fontName="DejaVuSans-Bold")
    normal = ParagraphStyle("normalc", parent=styles["Normal"], fontName="DejaVuSans")

    elems = [
        Paragraph("VISITA A CLIENTES", h1),
        Paragraph("CMAC Caja Arequipa — Unidad de Auditoría Interna", normal),
        Paragraph(f"Auditor: {usuario} · Fecha de visita: {ahora_peru().strftime('%d/%m/%Y %H:%M')} (hora Perú)", normal),
        Spacer(1, 10),
    ]

    ancho_contenido = docpdf.width

    numero = 1

    if criterios_txt:
        elems.extend(_pdf_seccion_lista(f"{_romano(numero)}. Criterio para la visita", ancho_contenido, criterios_txt))
        if observacion_criterio:
            elems.extend(_pdf_seccion_parrafo(f"{_romano(numero)}.1 Observación", ancho_contenido, observacion_criterio))
        numero += 1
    elif observacion_criterio:
        elems.extend(_pdf_seccion_parrafo(f"{_romano(numero)}. Observación", ancho_contenido, observacion_criterio))
        numero += 1

    elems.extend(_pdf_seccion_kv(f"{_romano(numero)}. Datos del cliente y crédito", ancho_contenido, [
        ("Agencia", safe_str(cliente.get("AGENCIA"))),
        ("DNI/LE Titular", safe_str(cliente.get("DOCPEN"))),
        ("Titular", safe_str(cliente.get("CLIENTE"))),
        ("Cuenta cliente", safe_str(cliente.get("BCCTA"))),
        ("N° de operación", safe_str(cliente.get("BCOPER"))),
        ("Módulo", safe_str(cliente.get("MODULO"))),
        ("Analista vigente", safe_str(cliente.get("ANALISTA"))),
        ("Analista evaluador", safe_str(cliente.get("ANALISTA_EVAL"))),
        ("Auditor (visita)", usuario),
        ("Importe", fmt_money(cliente.get("IMPDESEMB_MN"))),
        ("Saldo capital", fmt_money(cliente.get("SALDO_MN"))),
        ("Tipo de crédito", safe_str(cliente.get("PRODUCTO_CAJA"))),
        ("Tipo SBS", safe_str(cliente.get("TIPO_SBS"))),
        ("Calificación", safe_str(cliente.get("CATEG_RESULTANTE"))),
        ("Rubro", safe_str(cliente.get("ACTIVIDAD_ECON"))),
        ("Último pago", safe_str(cliente.get("FECHA_UTLPAGO"))),
        ("Resultado de la visita / Cliente visitado", cliente_visitado or "-"),
    ]))
    numero += 1

    for clave, etiqueta in [("negocio", "Negocio"),
                             ("laboral", "Laboral"),
                             ("aval", "Aval"),
                             ("domicilio", "Domicilio")]:
        d = visitas.get(clave)
        if not d:
            continue  # solo se incluyen las visitas con información registrada
        elems.extend(build_visita_section_pdf(_romano(numero), clave, etiqueta, d, cliente_visitado, ancho_contenido))
        numero += 1

    if garantias:
        elems.extend(_pdf_seccion_grupos(f"{_romano(numero)}. Garantías", ancho_contenido,
                                          [list(g.items()) for g in garantias], "Garantía"))
        numero += 1

    if rcc:
        elems.extend(_pdf_seccion_grupos(f"{_romano(numero)}. Deuda RCC", ancho_contenido,
                                          [list(r.items()) for r in rcc], "Deuda"))
        numero += 1

    elems.extend(_pdf_seccion_kv(f"{_romano(numero)}. Conformidad", ancho_contenido, [
        ("Hecho por (Auditor)", usuario), ("Fecha", ahora_peru().strftime("%d/%m/%Y")),
    ], dos_columnas=False))

    docpdf.build(elems)
    buf.seek(0)
    return buf
    
def to_excel(df):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='Resumen')
        return output.getvalue()
